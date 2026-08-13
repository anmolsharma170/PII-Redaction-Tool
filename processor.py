# pyrefly: ignore [missing-import]
import os
from docx import Document
from redactor import redact_text_content

def process_prospectus_docx(input_path: str, output_path: str):
    """
    Reads a .docx file, redacts PII from both standard text paragraphs 
    and tables, and saves the redacted version to a new file.
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")
        
    doc = Document(input_path)

    # Process standard text paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            para.text = redact_text_content(para.text)

    # Process text inside table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        para.text = redact_text_content(para.text)

    # Ensure output directory exists
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    doc.save(output_path)
    print(f"Document successfully redacted and saved to {output_path}")

if __name__ == "__main__":
    # Create a dummy .docx document to test the pipeline locally
    test_input = "temp_test_prospectus.docx"
    test_output = "temp_redacted_prospectus.docx"
    
    print(f"Creating a temporary docx file '{test_input}' to test processing...")
    doc = Document()
    doc.add_heading("Red Herring Prospectus Draft", 0)
    
    doc.add_paragraph("This is a confidential draft prospectus. Prepared by Jane Doe at Acme Corp.")
    doc.add_paragraph("For enquiries, email contact@acmecorp.com or call 555-987-6543.")
    
    table = doc.add_table(rows=2, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Employee Name"
    hdr_cells[1].text = "Contact Details"
    
    row_cells = table.rows[1].cells
    row_cells[0].text = "Alice Smith"
    row_cells[1].text = "alice.smith@acmecorp.com"
    
    doc.save(test_input)
    
    print("Processing the document for redaction...")
    process_prospectus_docx(test_input, test_output)
    
    # Read the output to verify
    redacted_doc = Document(test_output)
    print("\n--- Redacted Document Content ---")
    for para in redacted_doc.paragraphs:
        if para.text.strip():
            print(f"P: {para.text}")
            
    for tbl in redacted_doc.tables:
        for row in tbl.rows:
            row_texts = [cell.text for cell in row.cells]
            print(f"T: {row_texts}")
            
    # Clean up test files
    try:
        os.remove(test_input)
        os.remove(test_output)
        print("Cleaned up temporary test files.")
    except Exception as e:
        print(f"Error cleaning up: {e}")
